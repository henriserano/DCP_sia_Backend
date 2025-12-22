from __future__ import annotations

from app.extractors.base import BaseExtractor, ExtractResult, normalize_path


class DocxExtractor:
    name = "docx"
    supported_suffixes = {".docx"}

    def can_extract(self, path: str) -> bool:
        p = normalize_path(path)
        return p.is_file() and p.suffix.lower() in self.supported_suffixes

    def extract(self, path: str) -> ExtractResult:
        p = normalize_path(path)
        try:
            import docx
        except Exception as e:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx") from e

        d = docx.Document(str(p))
        text = "\n".join([para.text for para in d.paragraphs if para.text])

        return ExtractResult(
            text=text.strip(),
            source_path=str(p),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={"extractor": self.name, "paragraphs": str(len(d.paragraphs))},
        )